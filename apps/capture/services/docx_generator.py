"""Word document generation service for Capture entries."""

import logging
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def _parse_markdown_to_docx(document, markdown_text):
    """
    Parse markdown text and add properly formatted content to a Word document.

    Handles:
    - ## Section headers
    - **bold** text
    - - bullet lists
    - Regular paragraphs

    Args:
        document: python-docx Document instance
        markdown_text: Markdown-formatted text string
    """
    if not markdown_text:
        return

    lines = markdown_text.split('\n')
    current_list_items = []
    in_list = False

    def flush_list():
        """Add accumulated list items to document."""
        nonlocal current_list_items, in_list
        if current_list_items:
            for item in current_list_items:
                p = document.add_paragraph(style='List Bullet')
                _add_formatted_text(p, item)
            current_list_items = []
        in_list = False

    for line in lines:
        stripped = line.strip()

        # Skip empty lines (but flush lists first)
        if not stripped:
            flush_list()
            continue

        # Section headers (## Header)
        if stripped.startswith('## '):
            flush_list()
            header_text = stripped[3:].strip()
            heading = document.add_heading(header_text, level=3)
            # Style the heading
            for run in heading.runs:
                run.font.color.rgb = RGBColor(37, 99, 235)  # Blue
                run.font.size = Pt(12)

        # Bullet list items (- item or * item)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            in_list = True
            item_text = stripped[2:].strip()
            current_list_items.append(item_text)

        # Regular paragraph or bold header line
        else:
            flush_list()
            p = document.add_paragraph()
            _add_formatted_text(p, stripped)

    # Flush any remaining list items
    flush_list()


def _add_formatted_text(paragraph, text):
    """
    Add text to a paragraph with proper formatting for bold markers.

    Handles **bold** markers within the text.

    Args:
        paragraph: python-docx Paragraph instance
        text: Text that may contain **bold** markers
    """
    # Pattern to match **bold** text
    pattern = r'\*\*([^*]+)\*\*'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        # Add bold text
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True

        last_end = match.end()

    # Add remaining text after last match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def generate_docx(capture_entry):
    """
    Generate a Word document for a capture entry.

    Args:
        capture_entry: CaptureEntry model instance

    Returns:
        bytes: DOCX file content as bytes
    """
    document = Document()

    # Set up styles
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Add header
    header = document.add_paragraph()
    header_run = header.add_run('Whole Life Journey')
    header_run.bold = True
    header_run.font.size = Pt(14)
    header_run.font.color.rgb = RGBColor(99, 102, 241)  # Indigo color
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run('Capture Summary')
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)  # Gray
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add horizontal line (using a paragraph with bottom border)
    document.add_paragraph()

    # Title
    title = document.add_heading(capture_entry.title or 'Untitled Recording', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata section
    meta_para = document.add_paragraph()

    # Date
    meta_para.add_run('Date: ').bold = True
    meta_para.add_run(capture_entry.created_at.strftime('%A, %B %d, %Y'))
    meta_para.add_run('\n')

    # Duration
    if capture_entry.duration_seconds:
        minutes = capture_entry.duration_seconds // 60
        seconds = capture_entry.duration_seconds % 60
        meta_para.add_run('Duration: ').bold = True
        meta_para.add_run(f'{minutes}:{seconds:02d}')
        meta_para.add_run('\n')

    # Category
    if capture_entry.category:
        meta_para.add_run('Category: ').bold = True
        category_text = capture_entry.get_category_display()
        if capture_entry.subcategory:
            category_text += f' / {capture_entry.get_subcategory_display()}'
        meta_para.add_run(category_text)

    document.add_paragraph()  # Spacing

    # Summary section
    if capture_entry.summary:
        # Parse and add markdown-formatted summary
        _parse_markdown_to_docx(document, capture_entry.summary)
        document.add_paragraph()  # Spacing

    # Transcript section
    if capture_entry.transcript:
        document.add_heading('Full Transcript', level=2)

        transcript_para = document.add_paragraph(capture_entry.transcript)
        transcript_para.style.font.size = Pt(10)
        transcript_para.style.font.color.rgb = RGBColor(107, 114, 128)

    # Footer
    document.add_paragraph()
    footer = document.add_paragraph()
    footer_run = footer.add_run('Generated from wholelifejourney.com')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(107, 114, 128)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_footer = document.add_paragraph()
    date_run = date_footer.add_run(
        f'{capture_entry.created_at.strftime("%B %d, %Y")} at {capture_entry.created_at.strftime("%I:%M %p")}'
    )
    date_run.font.size = Pt(8)
    date_run.font.color.rgb = RGBColor(156, 163, 175)
    date_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()
    docx_buffer.close()

    logger.info(
        "Generated DOCX for capture entry %s (%d bytes)",
        capture_entry.id,
        len(docx_bytes)
    )

    return docx_bytes


def get_docx_filename(capture_entry):
    """
    Generate a filename for the DOCX download.

    Args:
        capture_entry: CaptureEntry model instance

    Returns:
        str: Sanitized filename for the DOCX
    """
    # Use title or fallback to 'Capture'
    title = capture_entry.title or 'Capture'

    # Sanitize title for filename
    safe_title = "".join(
        c if c.isalnum() or c in ' -_' else '_'
        for c in title
    )

    # Truncate if too long
    if len(safe_title) > 50:
        safe_title = safe_title[:50]

    # Remove leading/trailing spaces and underscores
    safe_title = safe_title.strip(' _')

    # Format: Title - WLJ Capture - Date.docx
    date_str = capture_entry.created_at.strftime('%Y-%m-%d')

    return f"{safe_title} - WLJ Capture - {date_str}.docx"
